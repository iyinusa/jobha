import logging
from typing import Dict, List, Any
import docx

logger = logging.getLogger(__name__)

class DocxParser:
    """
    Parser for Word documents using python-docx
    """
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from a Word document
        
        Args:
            file_path (str): Path to the Word document
            
        Returns:
            str: Extracted text content
        """
        try:
            # Load the document
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs with proper spacing
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Join paragraphs with newlines
            full_text = '\n'.join(paragraphs)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            # Add table text if any
            if tables_text:
                full_text += '\n\n' + '\n'.join(tables_text)
                
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse Word document: {str(e)}")